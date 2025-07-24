import streamlit as st
import pandas as pd
import io
import re
from fractions import Fraction
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF

st.set_page_config(page_title="Share Text Parser & Exporter", layout="wide")
st.title("📋 Share Text Parser with Export to Excel, Word & PDF")

def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def df_to_word(df):
    doc = Document()
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        run = hdr_cells[i].paragraphs[0].add_run(str(col))
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 128)
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
            if idx % 2 == 1:
                set_cell_background(row_cells[i], "D3D3D3")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Share Data', border=0, ln=1, align='C')
        self.ln(5)

def df_to_pdf(df):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    col_width = pdf.w / (len(df.columns) + 1)
    pdf.set_font("Arial", 'B', 11)
    for col in df.columns:
        pdf.cell(col_width, 10, str(col), border=1, align='C')
    pdf.ln()
    pdf.set_font("Arial", size=10)
    for _, row in df.iterrows():
        for item in row:
            pdf.cell(col_width, 10, str(item), border=1)
        pdf.ln()
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

def parse_share_text(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    owners = []
    current_owner_lines = []
    current_fraction = ""

    def flush_owner():
        if current_owner_lines and current_fraction:
            # Join owner lines as name + narration
            full_name = " ".join(current_owner_lines[:-1]) + (" " + current_owner_lines[-1] if len(current_owner_lines) > 1 else "")
            owners.append((full_name.strip(), current_fraction))

    for line in lines:
        # Check if line contains fraction (match digits/digits भाग)
        if re.search(r'\d+/\d+\s*भाग', line):
            # Flush previous owner if any
            if current_owner_lines:
                current_fraction = line
                # Owner name is everything before fraction line, so flush now
                flush_owner()
                current_owner_lines = []
                current_fraction = ""
        else:
            # Collect owner name / narration line
            current_owner_lines.append(line)
    # For any leftover
    if current_owner_lines:
        # No fraction follows, assign empty fraction
        owners.append((" ".join(current_owner_lines), ""))
    return owners

st.subheader("Paste Share Text Here")
input_text = st.text_area("Paste the unstructured share text data here:")

if st.button("Parse and Show Data"):
    parsed_data = parse_share_text(input_text)
    if not parsed_data:
        st.warning("No valid data found.")
    else:
        df = pd.DataFrame(parsed_data, columns=["Owner Name + Narration", "Share Fraction"])

        # Convert fraction to decimal
        def frac_to_decimal(frac_str):
            try:
                frac_part = frac_str.split()[0]  # e.g. '603/3076'
                return float(Fraction(frac_part))
            except:
                return None

        df["Share (Decimal)"] = df["Share Fraction"].apply(frac_to_decimal)

        st.dataframe(df)

        # Prepare Excel export
        output_excel = io.BytesIO()
        with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Share Data")
        output_excel.seek(0)

        st.download_button("Download Excel File", output_excel,
                           file_name="share_data.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # Prepare Word export
        word_file = df_to_word(df)
        st.download_button("Download Word File", word_file,
                           file_name="share_data.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Prepare PDF export
        pdf_file = df_to_pdf(df)
        st.download_button("Download PDF File", pdf_file,
                           file_name="share_data.pdf",
                           mime="application/pdf")

